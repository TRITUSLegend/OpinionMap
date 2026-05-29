from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
import io
import re
from app.models.report import Report
from app.core.logging import get_logger

logger = get_logger(__name__)

async def get_report(db: AsyncSession, report_id) -> Report | None:
    result = await db.execute(select(Report).where(Report.id == report_id))
    return result.scalar_one_or_none()

async def list_reports(db: AsyncSession, user_id, skip: int = 0, limit: int = 100) -> tuple[list[Report], int]:
    result = await db.execute(
        select(Report)
        .where(Report.user_id == user_id)
        .order_by(Report.created_at.desc())
        .offset(skip)
        .limit(limit)
    )
    reports = list(result.scalars().all())
    
    count_result = await db.execute(select(func.count(Report.id)).where(Report.user_id == user_id))
    total = count_result.scalar()
    
    return reports, total

from sqlalchemy import delete
from app.models.workflow import Workflow

async def delete_report(db: AsyncSession, report_id) -> bool:
    report = await get_report(db, report_id)
    if not report:
        return False
        
    workflow_id = report.workflow_id
    
    # First delete the report to avoid foreign key constraints
    result = await db.execute(delete(Report).where(Report.id == report_id))
    
    # Then cascade delete the workflow if it existed
    if workflow_id:
        from app.models.scraped_data import ScrapedData
        from app.models.agent_log import AgentLog
        from app.models.analytics import Analytics
        from app.models.embedding_metadata import EmbeddingMetadata
        
        await db.execute(delete(EmbeddingMetadata).where(EmbeddingMetadata.workflow_id == workflow_id))
        await db.execute(delete(Analytics).where(Analytics.workflow_id == workflow_id))
        await db.execute(delete(AgentLog).where(AgentLog.workflow_id == workflow_id))
        await db.execute(delete(ScrapedData).where(ScrapedData.workflow_id == workflow_id))
        await db.execute(delete(Workflow).where(Workflow.id == workflow_id))
        
    await db.commit()
    return result.rowcount > 0


def _sanitize_html(text: str) -> str:
    """Clean text for ReportLab XML parser — strip markdown and escape special chars."""
    if not text:
        return ""
    # Strip markdown bold
    text = text.replace('**', '')
    text = text.replace('*', '')
    # Replace common Unicode characters that ReportLab standard fonts don't support
    text = text.replace('\u20b9', 'Rs. ') # Rupee
    text = text.replace('\u2013', '-')    # En dash
    text = text.replace('\u2014', '--')   # Em dash
    text = text.replace('\u2018', "'")    # Left single quote
    text = text.replace('\u2019', "'")    # Right single quote
    text = text.replace('\u201c', '"')    # Left double quote
    text = text.replace('\u201d', '"')    # Right double quote
    text = text.replace('\u2026', '...')  # Ellipsis
    # Strip any remaining characters that are outside latin-1 (ReportLab's Helvetica only supports Latin-1)
    text = text.encode('latin1', 'ignore').decode('latin1')
    
    # Escape ampersands that aren't already escaped
    text = text.replace('&', '&amp;')
    # Escape angle brackets that aren't our own tags
    text = re.sub(r'<(?!/?(b|i|u|font|br|para|super|sub)[ >/])', '&lt;', text)
    return text


def generate_pdf(report: Report) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import (
            BaseDocTemplate, PageTemplate, Frame, Paragraph,
            Spacer, PageBreak, Table, TableStyle, HRFlowable
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        import datetime
        import json
        
        buffer = io.BytesIO()
        W, H = letter  # 612 x 792
        
        # ── Color Palette ──────────────────────────────────
        C_PRIMARY   = colors.HexColor('#4F46E5')  # Indigo 600
        C_SECONDARY = colors.HexColor('#0EA5E9')  # Sky 500
        C_DARK      = colors.HexColor('#0F172A')  # Slate 900
        C_TEXT      = colors.HexColor('#334155')  # Slate 700
        C_MUTED     = colors.HexColor('#64748B')  # Slate 500
        C_ACCENT    = colors.HexColor('#F43F5E')  # Rose 500
        C_SUCCESS   = colors.HexColor('#10B981')  # Emerald 500
        C_WARN      = colors.HexColor('#F59E0B')  # Amber 500
        C_BG_LIGHT  = colors.HexColor('#F1F5F9')  # Slate 100
        C_BG_CARD   = colors.HexColor('#EEF2FF')  # Indigo 50
        C_WHITE     = colors.white

        # ── Page Chrome ────────────────────────────────────
        def _page_chrome(canvas_obj, doc):
            canvas_obj.saveState()
            # Header bar
            canvas_obj.setFillColor(C_PRIMARY)
            canvas_obj.rect(0, H - 36, W, 36, stroke=0, fill=1)
            canvas_obj.setFillColor(C_WHITE)
            canvas_obj.setFont('Helvetica-Bold', 10)
            canvas_obj.drawString(40, H - 26, 'OpinionMap — Research & Analysis Report')
            canvas_obj.setFont('Helvetica', 9)
            canvas_obj.drawRightString(W - 40, H - 26,
                datetime.datetime.now().strftime('%B %d, %Y'))
            # Footer bar
            canvas_obj.setFillColor(C_DARK)
            canvas_obj.rect(0, 0, W, 28, stroke=0, fill=1)
            canvas_obj.setFillColor(C_WHITE)
            canvas_obj.setFont('Helvetica', 8)
            canvas_obj.drawString(40, 9, 'Confidential — Generated by OpinionMap')
            canvas_obj.drawRightString(W - 40, 9, f'Page {doc.page}')
            # Thin accent line under header
            canvas_obj.setStrokeColor(C_SECONDARY)
            canvas_obj.setLineWidth(2)
            canvas_obj.line(0, H - 37, W, H - 37)
            canvas_obj.restoreState()

        doc = BaseDocTemplate(buffer, pagesize=letter,
            leftMargin=50, rightMargin=50, topMargin=52, bottomMargin=42)
        frame = Frame(doc.leftMargin, doc.bottomMargin,
                      doc.width, doc.height, id='main')
        doc.addPageTemplates([
            PageTemplate(id='standard', frames=frame, onPage=_page_chrome)
        ])

        # ── Styles ─────────────────────────────────────────
        ss = getSampleStyleSheet()
        
        s_title = ParagraphStyle('RTitle', parent=ss['Heading1'],
            fontSize=30, leading=36, textColor=C_PRIMARY,
            alignment=TA_CENTER, fontName='Helvetica-Bold',
            spaceBefore=0, spaceAfter=6)
        
        s_tagline = ParagraphStyle('RTagline', parent=ss['Normal'],
            fontSize=14, leading=20, textColor=C_SECONDARY,
            alignment=TA_CENTER, fontName='Helvetica-Oblique',
            spaceAfter=60)
        
        s_cover_meta = ParagraphStyle('RCoverMeta', parent=ss['Normal'],
            fontSize=11, leading=16, textColor=C_MUTED,
            alignment=TA_CENTER, fontName='Helvetica')
        
        s_section = ParagraphStyle('RSectionHead', parent=ss['Heading2'],
            fontSize=16, leading=22, textColor=C_PRIMARY,
            fontName='Helvetica-Bold', spaceBefore=22, spaceAfter=10,
            borderColor=C_PRIMARY, borderWidth=0, borderPadding=0)
        
        s_body = ParagraphStyle('RBody', parent=ss['Normal'],
            fontSize=10.5, leading=17, textColor=C_TEXT,
            fontName='Helvetica', spaceAfter=10,
            firstLineIndent=0)
        
        s_bullet = ParagraphStyle('RBullet', parent=ss['Normal'],
            fontSize=10.5, leading=16, textColor=C_TEXT,
            fontName='Helvetica', leftIndent=20, spaceAfter=5,
            bulletIndent=8)
        
        s_metric_label = ParagraphStyle('RMetricLabel', parent=ss['Normal'],
            fontSize=10, leading=14, textColor=C_DARK,
            fontName='Helvetica-Bold')
        
        s_metric_val = ParagraphStyle('RMetricVal', parent=ss['Normal'],
            fontSize=10, leading=14, textColor=C_TEXT,
            fontName='Helvetica')

        # ── Data Extraction ────────────────────────────────
        sections = {}
        tagline = ''
        adv_metrics = {}
        recommendations = report.recommendations or []
        
        try:
            if report.full_report:
                raw = json.loads(report.full_report)
                # Handle both old format (sections-only) and new format (full dict)
                if 'sections' in raw:
                    sections = raw.get('sections', {})
                    tagline = raw.get('tagline', '')
                    adv_metrics = raw.get('advanced_metrics', {})
                    recommendations = raw.get('recommendations', recommendations)
                else:
                    # Old format: raw IS the sections dict
                    sections = raw
        except Exception:
            pass

        flow = []  # master flowables list

        # ━━━━━━━━━━━━━━━━ COVER PAGE ━━━━━━━━━━━━━━━━━━━━
        flow.append(Spacer(1, 140))
        flow.append(Paragraph(_sanitize_html(report.title or 'Product Intelligence Report'), s_title))
        if tagline:
            flow.append(Spacer(1, 8))
            flow.append(Paragraph(f'"{_sanitize_html(tagline)}"', s_tagline))
        else:
            flow.append(Spacer(1, 60))
        
        flow.append(Spacer(1, 40))
        flow.append(HRFlowable(width='40%', thickness=2, color=C_SECONDARY,
                                spaceAfter=20, spaceBefore=0, hAlign='CENTER'))
        flow.append(Paragraph('Generated by <b>OpinionMap</b>', s_cover_meta))
        flow.append(PageBreak())

        # ━━━━━━━━ HELPER: render a section body ━━━━━━━━━
        def _render_body(text: str):
            """Split text into paragraphs and render with bullets where appropriate."""
            if not text:
                return
            paragraphs = str(text).split('\n')
            for p in paragraphs:
                p = p.strip()
                if not p:
                    continue
                clean = _sanitize_html(p)
                # Detect bullet lines
                if re.match(r'^[-•■►]\s', p) or re.match(r'^\d+[.)]\s', p):
                    # Strip leading marker
                    bullet_text = re.sub(r'^[-•■►]\s*', '', clean)
                    bullet_text = re.sub(r'^\d+[.)]\s*', '', bullet_text)
                    flow.append(Paragraph(
                        f'<font color="#0EA5E9">▸</font>  {bullet_text}', s_bullet))
                else:
                    flow.append(Paragraph(clean, s_body))

        # ━━━━━━━━━━ EXECUTIVE SUMMARY ━━━━━━━━━━━━━━━━━━━
        flow.append(Paragraph('Executive Summary', s_section))
        flow.append(HRFlowable(width='100%', thickness=1, color=C_BG_LIGHT,
                                spaceAfter=10))
        _render_body(report.executive_summary)

        # ━━━━━━━━━━ ADVANCED METRICS SCORECARD ━━━━━━━━━━
        if adv_metrics:
            flow.append(Spacer(1, 6))
            flow.append(Paragraph('Product Evaluation Scorecard', s_section))
            flow.append(HRFlowable(width='100%', thickness=1, color=C_BG_LIGHT,
                                    spaceAfter=8))
            # Build a table of metrics
            metric_data = [
                [Paragraph('<b>Metric</b>', s_metric_label),
                 Paragraph('<b>Score / Rating</b>', s_metric_label)]
            ]
            for k, v in adv_metrics.items():
                metric_data.append([
                    Paragraph(_sanitize_html(k), s_metric_label),
                    Paragraph(_sanitize_html(str(v)), s_metric_val)
                ])
            t = Table(metric_data, colWidths=[2.4*inch, 4.1*inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), C_PRIMARY),
                ('TEXTCOLOR', (0, 0), (-1, 0), C_WHITE),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), C_BG_CARD),
                ('GRID', (0, 0), (-1, -1), 0.5, C_BG_LIGHT),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [C_WHITE, C_BG_CARD]),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ]))
            flow.append(t)
            flow.append(Spacer(1, 10))

        # ━━━━━━━━━━ ALL SECTIONS ━━━━━━━━━━━━━━━━━━━━━━━━
        for sec_title, content in sections.items():
            flow.append(Paragraph(_sanitize_html(sec_title), s_section))
            flow.append(HRFlowable(width='100%', thickness=1,
                                    color=C_BG_LIGHT, spaceAfter=8))
            
            if isinstance(content, list):
                for item in content:
                    flow.append(Paragraph(
                        f'<font color="#0EA5E9">▸</font>  {_sanitize_html(str(item))}',
                        s_bullet))
            elif isinstance(content, dict):
                for k, v in content.items():
                    if isinstance(v, list):
                        val_str = ', '.join(str(x) for x in v)
                        flow.append(Paragraph(
                            f'<b><font color="#4F46E5">{_sanitize_html(k)}:</font></b> {_sanitize_html(val_str)}',
                            s_body))
                    else:
                        flow.append(Paragraph(
                            f'<b><font color="#4F46E5">{_sanitize_html(k)}:</font></b> {_sanitize_html(str(v))}',
                            s_body))
            else:
                _render_body(str(content))

        # ━━━━━━━━━━ RECOMMENDATIONS ━━━━━━━━━━━━━━━━━━━━━
        if recommendations:
            flow.append(Paragraph('Strategic Recommendations', s_section))
            flow.append(HRFlowable(width='100%', thickness=1,
                                    color=C_BG_LIGHT, spaceAfter=8))
            for idx, rec in enumerate(recommendations, 1):
                flow.append(Paragraph(
                    f'<font color="#F43F5E"><b>{idx}.</b></font>  {_sanitize_html(str(rec))}',
                    s_bullet))

        # ━━━━━━━━━━ DISCLAIMER FOOTER ━━━━━━━━━━━━━━━━━━━
        flow.append(Spacer(1, 30))
        flow.append(HRFlowable(width='100%', thickness=1, color=C_MUTED,
                                spaceAfter=8))
        disc_style = ParagraphStyle('Disclaimer', parent=ss['Normal'],
            fontSize=8, leading=11, textColor=C_MUTED,
            fontName='Helvetica-Oblique', alignment=TA_CENTER)
        flow.append(Paragraph(
            'This report was generated autonomously by the OpinionMap multi-agent pipeline. '
            'All insights are derived from publicly available data and AI-powered analysis. '
            'This document is intended for strategic planning purposes only.',
            disc_style))

        doc.build(flow)
        buffer.seek(0)
        return buffer.read()

    except Exception as e:
        logger.error(f"Failed to generate PDF: {e}")
        return b"%PDF-1.4 Error generating PDF"


def generate_docx(report: Report) -> bytes:
    try:
        import docx
        import json
        doc = docx.Document()
        
        doc.add_heading(report.title or "Market Research Report", 0)
        
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.executive_summary or "No summary available.")
        
        # Render sections from full_report
        sections = {}
        try:
            if report.full_report:
                raw = json.loads(report.full_report)
                sections = raw.get('sections', raw) if 'sections' in raw else raw
        except Exception:
            pass
        
        for title, content in sections.items():
            doc.add_heading(title, level=1)
            if isinstance(content, str):
                for p in content.split('\n'):
                    if p.strip():
                        doc.add_paragraph(p.strip())
            elif isinstance(content, list):
                for item in content:
                    doc.add_paragraph(str(item), style='List Bullet')
            elif isinstance(content, dict):
                for k, v in content.items():
                    doc.add_paragraph(f"{k}: {v}")
        
        doc.add_heading("Recommendations", level=1)
        for rec in (report.recommendations or []):
            doc.add_paragraph(rec, style='List Bullet')
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    except Exception as e:
        logger.error(f"Failed to generate DOCX: {e}")
        return b"Error generating DOCX"
